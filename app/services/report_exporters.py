"""Render a project report dict (see app/services/report_service.py) into a
downloadable file. Each exporter returns raw bytes; the router wraps them in
a StreamingResponse with the matching content type — nothing is written to
disk, so concurrent requests for different projects never collide.
"""
from datetime import date, datetime
from io import BytesIO


def _fmt_money(value) -> str:
    return f"${float(value or 0):,.4f}"


def _fmt_dt(value) -> str:
    if value is None:
        return "—"
    if isinstance(value, (datetime, date)):
        return value.strftime("%Y-%m-%d %H:%M") if isinstance(value, datetime) else value.isoformat()
    return str(value)


def _period_label(report: dict) -> str:
    period = report["period"]
    start, end = period.get("start"), period.get("end")
    if not start and not end:
        return "All time"
    return f"{start or '…'} to {end or '…'}"


def _report_title(report: dict) -> str:
    return f"Governance Report — {report['project']['name']}"


# ---------------------------------------------------------------------------
# Excel (.xlsx) — openpyxl, already a project dependency
# ---------------------------------------------------------------------------

def export_excel(report: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    header_fill = PatternFill(start_color="9B1C5A", end_color="9B1C5A", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    title_font = Font(bold=True, size=14)

    def write_table(ws, start_row: int, headers: list[str], rows: list[list]):
        for col, h in enumerate(headers, start=1):
            cell = ws.cell(row=start_row, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for r, row in enumerate(rows, start=start_row + 1):
            for c, value in enumerate(row, start=1):
                ws.cell(row=r, column=c, value=value)
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 22
        return start_row + len(rows) + 2

    # ---- Overview sheet ----
    ws = wb.active
    ws.title = "Overview"
    ws["A1"] = _report_title(report)
    ws["A1"].font = title_font
    ws["A2"] = f"Organization: {report['project']['org_name']} ({report['project']['org_id']})"
    ws["A3"] = f"Project: {report['project']['name']} ({report['project']['id']})"
    ws["A4"] = f"Environment: {report['project']['environment'] or '—'}"
    ws["A5"] = f"Period: {_period_label(report)}"
    ws["A6"] = f"Generated: {_fmt_dt(report['generated_at'])} UTC"

    s = report["summary"]
    row = write_table(ws, 8, ["Metric", "Value"], [
        ["Total requests", s["total_requests"]],
        ["Input tokens", s["input_tokens"]],
        ["Output tokens", s["output_tokens"]],
        ["Total tokens", s["total_tokens"]],
        ["Input cost", _fmt_money(s["input_cost"])],
        ["Output cost", _fmt_money(s["output_cost"])],
        ["Total cost", _fmt_money(s["total_cost"])],
    ])

    # ---- Cost by model sheet ----
    ws2 = wb.create_sheet("Cost by Model")
    write_table(
        ws2, 1,
        ["Model", "Provider", "Requests", "Total Tokens", "Total Cost"],
        [
            [m["model_name"], m["provider"], m["total_requests"], m["total_tokens"], _fmt_money(m["total_cost"])]
            for m in report["cost_by_model"]
        ] or [["No usage in this period", "", "", "", ""]],
    )

    # ---- Budgets sheet ----
    ws3 = wb.create_sheet("Budgets")
    write_table(
        ws3, 1,
        ["Budget Type", "Limit", "Current Spend", "Utilization %", "Status"],
        [
            [
                b["budget_type"], _fmt_money(b["limit_amount"]) if b["limit_amount"] else "—",
                _fmt_money(b["current_spend"]), b["utilization_percent"], b["status"],
            ]
            for b in report["budgets"]
        ] or [["No budgets configured", "", "", "", ""]],
    )

    # ---- Governance rules sheet ----
    ws4 = wb.create_sheet("Governance Rules")
    write_table(
        ws4, 1,
        ["Rule Name", "Metric", "Operator", "Threshold", "Severity", "Scope"],
        [
            [r["rule_name"], r["metric_name"], r["operator"], r["threshold_value"], r["severity"], r["scope_level"]]
            for r in report["governance_rules"]
        ] or [["No active governance rules", "", "", "", "", ""]],
    )

    # ---- Alerts sheet ----
    ws5 = wb.create_sheet("Alerts")
    write_table(
        ws5, 1,
        ["Type", "Severity", "Status", "Message", "Occurred At"],
        [
            [a["alert_type"], a["severity"], a["status"], a["message"], _fmt_dt(a["created_at"])]
            for a in report["alerts"]
        ] or [["No alerts in this period", "", "", "", ""]],
    )

    # ---- Audit trail sheet ----
    ws6 = wb.create_sheet("Audit Trail")
    write_table(
        ws6, 1,
        ["Action", "Category", "Status", "Actor", "Summary", "Occurred At"],
        [
            [e["audit_action"], e["audit_category"], e["audit_status"], e["actor_id"] or "—",
             e["change_summary"] or "", _fmt_dt(e["occurred_at"])]
            for e in report["audit_entries"]
        ] or [["No audit entries in this period", "", "", "", "", ""]],
    )

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Word (.docx) — python-docx
# ---------------------------------------------------------------------------

def export_docx(report: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading(_report_title(report), level=0)
    title.runs[0].font.color.rgb = RGBColor(0x95, 0x13, 0x5A)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Organization: {report['project']['org_name']} ({report['project']['org_id']})\n"
        f"Project: {report['project']['name']} ({report['project']['id']})\n"
        f"Environment: {report['project']['environment'] or '—'}\n"
        f"Period: {_period_label(report)}\n"
        f"Generated: {_fmt_dt(report['generated_at'])} UTC"
    )

    def add_table(headers: list[str], rows: list[list], empty_msg: str):
        if not rows:
            doc.add_paragraph(empty_msg).italic = True
            return
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)

    doc.add_heading("Usage Summary", level=1)
    s = report["summary"]
    add_table(
        ["Metric", "Value"],
        [
            ["Total requests", s["total_requests"]],
            ["Input tokens", s["input_tokens"]],
            ["Output tokens", s["output_tokens"]],
            ["Total tokens", s["total_tokens"]],
            ["Input cost", _fmt_money(s["input_cost"])],
            ["Output cost", _fmt_money(s["output_cost"])],
            ["Total cost", _fmt_money(s["total_cost"])],
        ],
        "No usage data.",
    )

    doc.add_heading("Cost by Model", level=1)
    add_table(
        ["Model", "Provider", "Requests", "Total Tokens", "Total Cost"],
        [
            [m["model_name"], m["provider"], m["total_requests"], m["total_tokens"], _fmt_money(m["total_cost"])]
            for m in report["cost_by_model"]
        ],
        "No model usage in this period.",
    )

    doc.add_heading("Budgets", level=1)
    add_table(
        ["Budget Type", "Limit", "Current Spend", "Utilization %", "Status"],
        [
            [
                b["budget_type"], _fmt_money(b["limit_amount"]) if b["limit_amount"] else "—",
                _fmt_money(b["current_spend"]), f"{b['utilization_percent']}%", b["status"],
            ]
            for b in report["budgets"]
        ],
        "No budgets configured for this project.",
    )

    doc.add_heading("Governance Rules", level=1)
    add_table(
        ["Rule Name", "Metric", "Operator", "Threshold", "Severity", "Scope"],
        [
            [r["rule_name"], r["metric_name"], r["operator"], r["threshold_value"], r["severity"], r["scope_level"]]
            for r in report["governance_rules"]
        ],
        "No active governance rules apply to this project.",
    )

    doc.add_heading("Alerts", level=1)
    add_table(
        ["Type", "Severity", "Status", "Message", "Occurred At"],
        [
            [a["alert_type"], a["severity"], a["status"], a["message"], _fmt_dt(a["created_at"])]
            for a in report["alerts"]
        ],
        "No alerts in this period.",
    )

    doc.add_heading("Audit Trail", level=1)
    add_table(
        ["Action", "Category", "Status", "Actor", "Summary", "Occurred At"],
        [
            [e["audit_action"], e["audit_category"], e["audit_status"], e["actor_id"] or "—",
             e["change_summary"] or "", _fmt_dt(e["occurred_at"])]
            for e in report["audit_entries"]
        ],
        "No audit entries in this period.",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by the AI Governance Platform")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF — reportlab
# ---------------------------------------------------------------------------

def export_pdf(report: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    MAGENTA = colors.HexColor("#95135A")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Title"], textColor=MAGENTA, fontSize=20,
    )
    heading_style = ParagraphStyle(
        "SectionHeading", parent=styles["Heading2"], textColor=MAGENTA, spaceBefore=16,
    )
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9.5, leading=14)
    empty_style = ParagraphStyle("Empty", parent=styles["Italic"], textColor=colors.grey)

    def table(headers: list[str], rows: list[list], empty_msg: str, col_widths=None):
        if not rows:
            return Paragraph(empty_msg, empty_style)
        data = [headers] + [[str(v) for v in row] for row in rows]
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), MAGENTA),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DDDDDD")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F4FB")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
    )

    story = [
        Paragraph(_report_title(report), title_style),
        Spacer(1, 6),
        Paragraph(
            f"Organization: {report['project']['org_name']} ({report['project']['org_id']})<br/>"
            f"Project: {report['project']['name']} ({report['project']['id']})<br/>"
            f"Environment: {report['project']['environment'] or '—'}<br/>"
            f"Period: {_period_label(report)}<br/>"
            f"Generated: {_fmt_dt(report['generated_at'])} UTC",
            meta_style,
        ),
    ]

    s = report["summary"]
    story += [
        Paragraph("Usage Summary", heading_style),
        table(["Metric", "Value"], [
            ["Total requests", s["total_requests"]],
            ["Input tokens", s["input_tokens"]],
            ["Output tokens", s["output_tokens"]],
            ["Total tokens", s["total_tokens"]],
            ["Input cost", _fmt_money(s["input_cost"])],
            ["Output cost", _fmt_money(s["output_cost"])],
            ["Total cost", _fmt_money(s["total_cost"])],
        ], "No usage data.", col_widths=[2.5 * inch, 2.5 * inch]),

        Paragraph("Cost by Model", heading_style),
        table(
            ["Model", "Provider", "Requests", "Tokens", "Cost"],
            [
                [m["model_name"], m["provider"], m["total_requests"], m["total_tokens"], _fmt_money(m["total_cost"])]
                for m in report["cost_by_model"]
            ],
            "No model usage in this period.",
        ),

        Paragraph("Budgets", heading_style),
        table(
            ["Budget Type", "Limit", "Spend", "Utilization %", "Status"],
            [
                [
                    b["budget_type"], _fmt_money(b["limit_amount"]) if b["limit_amount"] else "—",
                    _fmt_money(b["current_spend"]), f"{b['utilization_percent']}%", b["status"],
                ]
                for b in report["budgets"]
            ],
            "No budgets configured for this project.",
        ),

        Paragraph("Governance Rules", heading_style),
        table(
            ["Rule Name", "Metric", "Operator", "Threshold", "Severity"],
            [
                [r["rule_name"], r["metric_name"], r["operator"], r["threshold_value"], r["severity"]]
                for r in report["governance_rules"]
            ],
            "No active governance rules apply to this project.",
        ),

        Paragraph("Alerts", heading_style),
        table(
            ["Type", "Severity", "Status", "Message"],
            [
                [a["alert_type"], a["severity"], a["status"], a["message"]]
                for a in report["alerts"]
            ],
            "No alerts in this period.",
        ),

        Paragraph("Audit Trail", heading_style),
        table(
            ["Action", "Category", "Status", "Actor", "Occurred At"],
            [
                [e["audit_action"], e["audit_category"], e["audit_status"], e["actor_id"] or "—", _fmt_dt(e["occurred_at"])]
                for e in report["audit_entries"]
            ],
            "No audit entries in this period.",
        ),
    ]

    doc.build(story)
    return buf.getvalue()
